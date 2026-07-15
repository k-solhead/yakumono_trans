import streamlit as st
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
from copy import deepcopy
from io import BytesIO
import os
import re
import unicodedata

def set_run_text_preserve_images(run, new_text):
    """run内の<w:drawing>等の画像要素を保持しつつ、テキスト部分のみを置き換える。

    通常の`run.text = ...`はrun内の全子要素を一度削除して<w:t>を追加するため、
    画像（<w:drawing>/<w:pict>）が失われる。ここでは<w:t>要素のテキストのみを更新する。
    """
    r = run._r
    t_elements = r.findall(qn('w:t'))
    if t_elements:
        t_elements[0].text = new_text
        t_elements[0].set(qn('xml:space'), 'preserve')
        for t in t_elements[1:]:
            t.text = ''
    elif new_text:
        t = OxmlElement('w:t')
        t.text = new_text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)

def zen_to_han(char):
    """全角英数記号を半角に変換"""
    cp = ord(char)
    # 全角英数記号 (FF01-FF5E) → 半角 (0021-007E)
    if 0xFF01 <= cp <= 0xFF5E:
        return chr(cp - 0xFEE0)
    # 全角スペース → 半角スペース
    if cp == 0x3000:
        return ' '
    return char

def normalize_url_email(text):
    """URL・メールアドレス内の全角文字を半角に変換"""
    # URL（全角文字混在対応）
    url_pattern = r'[\uFF48\uFF28hH][\uFF54\uFF34tT][\uFF54\uFF34tT][\uFF50\uFF30pP][\uFF53\uFF33sS]?[\uFF1A：:][\uFF0F／/][\uFF0F／/][^\s\u3000]+'
    # メールアドレス（全角文字混在対応）
    email_pattern = r'[\w\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19][\w.\-\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19\uFF0E\uFF0D]*[\uFF20@][\w\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19][\w.\-\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19\uFF0E\uFF0D]*'
    combined = f'({url_pattern})|({email_pattern})'
    def replace_match(m):
        return ''.join(zen_to_han(c) for c in m.group(0))
    return re.sub(combined, replace_match, text)

def apply_run_highlights(para):
    """ボールド・イタリック・ダッシュ文字ごとのハイライトを適用する。"""
    dash_highlights = {
        '—': WD_COLOR_INDEX.BLUE,
        '–': WD_COLOR_INDEX.VIOLET,
    }

    for run in list(para.runs):
        run_text = run.text
        if not run_text:
            continue

        current_highlight = run.font.highlight_color
        base_highlight = current_highlight
        if run.bold:
            base_highlight = WD_COLOR_INDEX.BRIGHT_GREEN
        if run.italic:
            base_highlight = WD_COLOR_INDEX.RED

        segments = []
        buffer = []
        segment_highlight = dash_highlights.get(run_text[0], base_highlight)

        for char in run_text:
            char_highlight = dash_highlights.get(char, base_highlight)
            if buffer and char_highlight != segment_highlight:
                segments.append((''.join(buffer), segment_highlight))
                buffer = [char]
                segment_highlight = char_highlight
                continue
            buffer.append(char)

        if buffer:
            segments.append((''.join(buffer), segment_highlight))

        if len(segments) == 1:
            if segments[0][1] != current_highlight:
                run.font.highlight_color = segments[0][1]
            continue

        is_bold = run.bold
        is_italic = run.italic
        original_rpr = run._r.find(qn('w:rPr'))
        rpr_template = deepcopy(original_rpr) if original_rpr is not None else None

        text0, highlight0 = segments[0]
        set_run_text_preserve_images(run, text0)
        if highlight0 != current_highlight:
            run.font.highlight_color = highlight0

        prev_r = run._r
        for seg_text, seg_highlight in segments[1:]:
            new_r = OxmlElement('w:r')
            if rpr_template is not None:
                new_r.append(deepcopy(rpr_template))
            t = OxmlElement('w:t')
            t.text = seg_text
            t.set(qn('xml:space'), 'preserve')
            new_r.append(t)
            prev_r.addnext(new_r)
            wrapped = Run(new_r, run._parent)
            wrapped.bold = is_bold
            wrapped.italic = is_italic
            if seg_highlight != current_highlight:
                wrapped.font.highlight_color = seg_highlight
            prev_r = new_r

output_word = "./output/output.docx"
os.makedirs(os.path.dirname(output_word), exist_ok=True)

replacement = {
    '\u3000':'\u0020',      # 全角空白を半角空白へ
    '\u30FB':'\u2022',      # 中黒をビュレットに変換
    '\uFF01':'\u0021',      # !
    '\uFF02':'\u0022',      # "
    '\uFF03':'\u0023',      # #
    '\uFF04':'\u0024',      # $
    '\uFF05':'\u0025',      # %
    '\uFF06':'\u0026',      # &
    '\uFF07':'\u0027',      # '
    '\uFF08':'\u0028',      # (
    '\uFF09':'\u0029',      # )
    '\uFF0A':'\u002A',      # *
    '\uFF0B':'\u002B',      # +
    '\uFF0C':'\u002C',      # ,
    '\uFF0D':'\u002D',      # -
    '\uFF0E':'\u002E',      # .
    '\uFF0F':'\u002F',      # /
    '\uFF10':'\u0030',      # 0
    '\uFF11':'\u0031',      # 1
    '\uFF12':'\u0032',      # 2
    '\uFF13':'\u0033',      # 3
    '\uFF14':'\u0034',      # 4
    '\uFF15':'\u0035',      # 5
    '\uFF16':'\u0036',      # 6
    '\uFF17':'\u0037',      # 7
    '\uFF18':'\u0038',      # 8
    '\uFF19':'\u0039',      # 9
    '\uFF1A':'\u003A',      # :
    '\uFF1B':'\u003B',      # ;
    '\uFF1C':'\u003C',      # <
    '\uFF1D':'\u003D',      # =
    '\uFF1E':'\u003E',      # >
    '\uFF1F':'\u003F',      # ?
    '\uFF20':'\u0040',      # @
    '\uFF21':'\u0041',      # A
    '\uFF22':'\u0042',      # B
    '\uFF23':'\u0043',      # C
    '\uFF24':'\u0044',      # D
    '\uFF25':'\u0045',      # E
    '\uFF26':'\u0046',      # F
    '\uFF27':'\u0047',      # G
    '\uFF28':'\u0048',      # H
    '\uFF29':'\u0049',      # I
    '\uFF2A':'\u004A',      # J
    '\uFF2B':'\u004B',      # K
    '\uFF2C':'\u004C',      # L
    '\uFF2D':'\u004D',      # M
    '\uFF2E':'\u004E',      # N
    '\uFF2F':'\u004F',      # O
    '\uFF30':'\u0050',      # P
    '\uFF31':'\u0051',      # Q
    '\uFF32':'\u0052',      # R
    '\uFF33':'\u0053',      # S
    '\uFF34':'\u0054',      # T
    '\uFF35':'\u0055',      # U
    '\uFF36':'\u0056',      # V
    '\uFF37':'\u0057',      # W
    '\uFF38':'\u0058',      # X
    '\uFF39':'\u0059',      # Y
    '\uFF3A':'\u005A',      # Z
    '\uFF3B':'\u005B',      # [
    '\uFF3C':'\u005C',      # \
    '\uFF3D':'\u005D',      # ]
    '\uFF3E':'\u005E',      # ^
    '\uFF3F':'\u005F',      # _
    '\uFF40':'\u0060',      # `
    '\uFF41':'\u0061',      # a
    '\uFF42':'\u0062',      # b
    '\uFF43':'\u0063',      # c
    '\uFF44':'\u0064',      # d
    '\uFF45':'\u0065',      # e
    '\uFF46':'\u0066',      # f
    '\uFF47':'\u0067',      # g
    '\uFF48':'\u0068',      # h
    '\uFF49':'\u0069',      # i
    '\uFF4A':'\u006A',      # j
    '\uFF4B':'\u006B',      # k
    '\uFF4C':'\u006C',      # l
    '\uFF4D':'\u006D',      # m
    '\uFF4E':'\u006E',      # n
    '\uFF4F':'\u006F',      # o
    '\uFF50':'\u0070',      # p
    '\uFF51':'\u0071',      # q
    '\uFF52':'\u0072',      # r
    '\uFF53':'\u0073',      # s
    '\uFF54':'\u0074',      # t
    '\uFF55':'\u0075',      # u
    '\uFF56':'\u0076',      # v
    '\uFF57':'\u0077',      # w
    '\uFF58':'\u0078',      # x
    '\uFF59':'\u0079',      # y
    '\uFF5A':'\u007A',      # z
    '\uFF5B':'\u007B',      # {
    '\uFF5C':'\u007C',      # |
    '\uFF5D':'\u007D',      # }
    '\uFF5E':'\u007E',      # ~
    '\u33A1':'m2',          # 機種依存文字（平方メートル）
    '\u33A5':'m3',          # 機種依存文字（立法メートル）
    '\\ ':'¥',
    '("':'("',
    '\\':'¥',
    "'":"'",
    '" ':'"\\s',
    ' "':'\\s"',
    '")':'”)',
    '".':'”.',
    '."':'.”',
    '・\t':'• ',
    ' )':')',
    ' . ':'. ',
    ' , ':', ',
    '" ':'” ',
    '( ':'(',
    "' ":"’ ",
    "'":"’",
    "• ":"•   ",
    '） ':'）',
    ' （':'（',
    '."':'.”',
    "' ":"’ ",
    ' 、':'、',
    '。 ':'。',
    ' 年':'年',
    ' 日':'日',
    ' 月':'月',
    '、 ':'、',
    '： ':'：',
    ' ：':'：',
    ' 円':'円',
    ' 時':'時',
    ' 分':'分',
    ' 百万円':'百万円',
    ' 億円':'億円',
    ' ％':'％',
    ' %':'%',
    ',"':',”',
    '  ':' ',
    '",':'”,',
    ' 」':'」',
    '－':'–',
    '（ ':'（',
    ' （':'（',
    '） ':'）',
    '( ':'(',
    ' "':' “',
    '」 ':'」',
    ' 「':'「',
    '］ ':'］',
    'Source：':'Source: ',
    '（Billions）':' (Billions) ',
    'Note：':'Note: ',
    '（Thousands）':'(Thousands) ',
    '（Millions）':' (Millions) ',
    '〜 ':'〜',
    '\t -':'\t\u2013',
    '¥-':'¥–',
    '$ ':'$',
    '➢':'>',
    '\u3000\t':'\t',
    '\\s\\':'\\s¥',
    '\\s\\s\\s':'\\s',
    '\t\\':'\t¥',
}

st.title("テキスト整形（英文）")
st.write("英文フォントの半角への修正や約物の自動変換をします")

option = st.radio("Word文書かテキストか整形対象を選択してください", ("Word文書", "テキスト文書"))

if option == "テキスト文書":
    if "text" not in st.session_state:
        st.session_state.text = ""
    st.session_state.text = st.text_area("テキストを貼りつけてください", height=300, placeholder="ここに貼りつけてください")
    
    if st.button("実行する"):    
        try:
            text = st.session_state.text
            text = normalize_url_email(text)
            for old,new in replacement.items():
                text = text.replace(old, new)
            text = re.sub(r'(?<!\s)\(', r' (', text)
            text = re.sub(r'\)(?!\s|.|,)', r') ', text)
            text = re.sub(r':(?![/\s])', r': ', text)
            text = re.sub(r'[ ]+(\r?\n)', r'\1', text)
        
            st.success("処理が完了しました。下記テキストをコピペしてください")
            with st.container(border=True):
                st.text(text)
            
        except Exception as e:
            st.error(f"ファイルの読み込み中にエラーが発生しました: {e}")
            
else:    
    # st.file_uploaderウィジェットの作成
    # type引数で受け付けるファイルの形式を'.docx'に限定
    uploaded_file = st.file_uploader("Wordファイル（.docx）をアップロード", type=['docx'])

    # ファイルがアップロードされた場合の処理
    if uploaded_file is not None:
        file_name = os.path.splitext(uploaded_file.name)[0]  # アップロードされたファイル名を取得
        st.success("ファイルが正常にアップロードされました。")

        # BytesIOでアップロードされたファイルを扱う
        doc_file = BytesIO(uploaded_file.getvalue())

        try:
            # python-docxでWordドキュメントを開く
            doc = docx.Document(doc_file)
            for para in doc.paragraphs:
                # Run単位の処理：テキスト置換 + 太字・斜体のハイライト（画像は保持）
                for run in para.runs:
                    t = run.text
                    t = normalize_url_email(t)
                    for old, new in replacement.items():
                        t = t.replace(old, new)
                    t = re.sub(r'(?<!\s)\(', r' (', t)
                    t = re.sub(r'\)(?!\s|.|,)', r') ', t)
                    t = re.sub(r':(?![/\s])', r': ', t)
                    set_run_text_preserve_images(run, t)

                apply_run_highlights(para)

                # 改行前の半角空白を削除（w:br直前のw:t末尾スペース除去）
                _prev_t = None
                for _elem in para._element.iter():
                    if _elem.tag == qn('w:t') and _elem.text:
                        _prev_t = _elem
                    elif _elem.tag == qn('w:br') and _prev_t is not None:
                        stripped = _prev_t.text.rstrip(' \t')
                        if stripped != _prev_t.text:
                            _prev_t.text = stripped

                # 箇条書きスタイルの解除とビュレット挿入
                is_list = para.style.name.startswith('List') or para._element.pPr is not None and para._element.pPr.find(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') is not None
                if is_list:
                    if para.runs:
                        set_run_text_preserve_images(para.runs[0], '• ' + para.runs[0].text)
                    # numPr（ナンバリング属性）を削除して箇条書き設定を解除
                    if para._element.pPr is not None:
                        numPr = para._element.pPr.find(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                        if numPr is not None:
                            para._element.pPr.remove(numPr)
                    para.style = doc.styles['Normal']

            # 変更を新しいPDFファイルに保存
            doc.save(output_word)
            st.success("処理が完了しました")

            # ドキュメントの内容を表示
            #st.subheader("処理されたWordファイルの内容")
        
            # ドキュメント内の各段落を読み込んで表示
            #for para in doc.paragraphs:
            #    st.write(para.text)
            st.success("ダウンロードボタンを押してください")
            with open(output_word, "rb") as file:
                word_data = file.read()
                # ダウンロードボタンを作成
                st.download_button(
                    label="Word文書をダウンロード",
                    data=word_data,
                    file_name=file_name+"_chk.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", # WordファイルのMIMEタイプ
                    on_click="ignore" # 再実行を無視する設定
                )
            print(f"ダウンロードしました。")

        
        except Exception as e:
            st.error(f"ファイルの読み込み中にエラーが発生しました: {e}")
    else:
        st.info("ファイルをアップロードしてください。")            
