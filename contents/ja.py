
import streamlit as st
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run
from copy import deepcopy
from io import BytesIO
import os
import re

def set_run_text_preserve_images(run, new_text):
    """run内の<w:drawing>等の画像要素を保持しつつ、テキスト部分のみを置き換える。

    通常の`run.text = ...`はrun内の全子要素を一度削除して<w:t>を追加するため、
    画像（<w:drawing>/<w:pict>）が失われる。ここでは<w:t>要素のテキストのみを更新する。
    """
    r = run._r
    t_elements = r.findall(qn('w:t'))
    control_tags = {qn('w:br'), qn('w:cr'), qn('w:tab')}
    has_controls = any(child.tag in control_tags for child in r)

    if t_elements and has_controls:
        text_index = 0
        cursor = 0

        for child in r:
            if child.tag == qn('w:t'):
                start = cursor
                while cursor < len(new_text) and new_text[cursor] not in '\n\r\t':
                    cursor += 1
                child.text = new_text[start:cursor]
                child.set(qn('xml:space'), 'preserve')
                text_index += 1
            elif child.tag in {qn('w:br'), qn('w:cr')}:
                if cursor < len(new_text) and new_text[cursor] == '\r':
                    cursor += 1
                if cursor < len(new_text) and new_text[cursor] == '\n':
                    cursor += 1
            elif child.tag == qn('w:tab'):
                if cursor < len(new_text) and new_text[cursor] == '\t':
                    cursor += 1

        for t in t_elements[text_index:]:
            t.text = ''
    elif t_elements:
        t_elements[0].text = new_text
        t_elements[0].set(qn('xml:space'), 'preserve')
        for t in t_elements[1:]:
            t.text = ''
    elif new_text:
        t = OxmlElement('w:t')
        t.text = new_text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)

def zen_to_han(char):
    """全角英数記号を半角に変換"""
    cp = ord(char)
    # 全角英数記号 (FF01-FF5E) → 半角 (0021-007E)
    if 0xFF01 <= cp <= 0xFF5E:
        return chr(cp - 0xFEE0)
    # 全角スペース → 半角スペース
    if cp == 0x3000:
        return ' '
    return char

def normalize_url_email(text):
    """URL・メールアドレス内の全角文字を半角に変換"""
    # URL（全角文字混在対応）
    url_pattern = r'[\uFF48\uFF28hH][\uFF54\uFF34tT][\uFF54\uFF34tT][\uFF50\uFF30pP][\uFF53\uFF33sS]?[\uFF1A：:][\uFF0F／/][\uFF0F／/][^\s\u3000]+'
    # メールアドレス（全角文字混在対応）
    email_pattern = r'[\w\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19][\w.\-\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19\uFF0E\uFF0D]*[\uFF20@][\w\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19][\w.\-\uFF21-\uFF3A\uFF41-\uFF5A\uFF10-\uFF19\uFF0E\uFF0D]*'
    combined = f'({url_pattern})|({email_pattern})'
    def replace_match(m):
        return ''.join(zen_to_han(c) for c in m.group(0))
    return re.sub(combined, replace_match, text)

# URLパターン（全角スラッシュ・コロン混在対応）
_url_re = re.compile(r'https?[:\uFF1A][/\uFF0F]{2}[^\s\u3000]+')

def protect_urls(text, func):
    """テキスト中のURLを一時的にプレースホルダーに置き、funcを適用後に復元。URL内の全角スラッシュ・コロンは半角に変換"""
    urls = []
    def save(m):
        url = m.group(0).replace('\uFF0F', '/').replace('\uFF1A', ':')
        urls.append(url)
        return f'\x00URL{len(urls)-1}\x00'
    text = _url_re.sub(save, text)
    text = func(text)
    for i, url in enumerate(urls):
        text = text.replace(f'\x00URL{i}\x00', url)
    return text

def apply_slash_colon(text):
    """/→／、:→：を変換（URL外）"""
    text = text.replace('/', '\uff0f')
    text = text.replace(':', '\uff1a')
    return text

def strip_line_head_spaces(text):
    """各行の先頭に混入した空白を除去する。"""
    return re.sub(r'(^|\r?\n)[ \u00A0\u3000\t]+', r'\1', text)

def strip_spaces_around_numbers(text):
    """数字前後の空白を除去しつつ、※/＊/*+数字の直後の半角スペースは保持する。"""
    protected_spaces = []

    def protect_note_space(match):
        protected_spaces.append(match.group(0))
        return f'\x00NOTE_SPACE{len(protected_spaces) - 1}\x00'

    text = re.sub(r'([\u203B\uFF0A\*][0-9\uFF10-\uFF19]+) ', protect_note_space, text)
    text = re.sub(r'\s*(\d+)\s*', r'\1', text)

    for index, protected in enumerate(protected_spaces):
        text = text.replace(f'\x00NOTE_SPACE{index}\x00', protected)

    return text

def apply_run_highlights(para):
    """ボールド・イタリック・ダッシュ文字ごとのハイライトを適用する。"""
    dash_highlights = {
        '—': WD_COLOR_INDEX.BLUE,
        '–': WD_COLOR_INDEX.VIOLET,
    }

    for run in list(para.runs):
        run_text = run.text
        if not run_text:
            continue

        current_highlight = run.font.highlight_color
        base_highlight = current_highlight
        if run.bold:
            base_highlight = WD_COLOR_INDEX.BRIGHT_GREEN
        if run.italic:
            base_highlight = WD_COLOR_INDEX.RED

        segments = []
        buffer = []
        segment_highlight = dash_highlights.get(run_text[0], base_highlight)

        for char in run_text:
            char_highlight = dash_highlights.get(char, base_highlight)
            if buffer and char_highlight != segment_highlight:
                segments.append((''.join(buffer), segment_highlight))
                buffer = [char]
                segment_highlight = char_highlight
                continue
            buffer.append(char)

        if buffer:
            segments.append((''.join(buffer), segment_highlight))

        if len(segments) == 1:
            if segments[0][1] != current_highlight:
                run.font.highlight_color = segments[0][1]
            continue

        is_bold = run.bold
        is_italic = run.italic
        original_rpr = run._r.find(qn('w:rPr'))
        rpr_template = deepcopy(original_rpr) if original_rpr is not None else None

        text0, highlight0 = segments[0]
        set_run_text_preserve_images(run, text0)
        if highlight0 != current_highlight:
            run.font.highlight_color = highlight0

        prev_r = run._r
        for seg_text, seg_highlight in segments[1:]:
            new_r = OxmlElement('w:r')
            if rpr_template is not None:
                new_r.append(deepcopy(rpr_template))
            t = OxmlElement('w:t')
            t.text = seg_text
            t.set(qn('xml:space'), 'preserve')
            new_r.append(t)
            prev_r.addnext(new_r)
            wrapped = Run(new_r, run._parent)
            wrapped.bold = is_bold
            wrapped.italic = is_italic
            if seg_highlight != current_highlight:
                wrapped.font.highlight_color = seg_highlight
            prev_r = new_r

# ※/＊/*+数字パターン（ハイライト対象）
note_pattern = re.compile(r'[\u203B\uFF0A\*][0-9\uFF10-\uFF19]+')

def apply_note_highlight_per_run(para):
    """段落内の※/＊/*+数字パターンをハイライトする。

    パターンがrunをまたぐ場合（数字部分がWordの上付き/下付きで別runになるなど）も対応するため、
    段落全体のテキストでマッチを取り、各runの該当部分のみをREDハイライトに切り出す。
    画像runの位置は保ち、ハイライトはマッチ部分のみに限定する（後続文字や次run・段落に波及させない）。
    """
    if not para.runs:
        return
    # 各runの段落全体テキスト中での位置を記録
    run_info = []  # (run, start, end)
    cursor = 0
    for run in para.runs:
        rt = run.text
        run_info.append((run, cursor, cursor + len(rt)))
        cursor += len(rt)
    full_text = ''.join(r.text for r, _, _ in run_info)

    matches = list(note_pattern.finditer(full_text))
    if not matches:
        return

    # 各runについて、自身に重なるマッチ部分（local座標）を集める
    for run, r_start, r_end in run_info:
        if r_start == r_end:
            continue  # 空run（画像runなど）はスキップ
        intersects = []
        for m in matches:
            i_start = max(m.start(), r_start)
            i_end = min(m.end(), r_end)
            if i_start < i_end:
                intersects.append((i_start - r_start, i_end - r_start))
        if not intersects:
            continue

        run_text = full_text[r_start:r_end]
        is_bold = run.bold
        is_italic = run.italic
        cur_highlight = run.font.highlight_color
        # rPrのテンプレートは「修正前」の状態を必ず deepcopy で確保する。
        # （後で run.font.highlight_color を上書きすると元のrPrが書き換わり、
        #  新run作成時にハイライトが波及してしまうため）
        original_rpr = run._r.find(qn('w:rPr'))
        rpr_template = deepcopy(original_rpr) if original_rpr is not None else None

        # セグメント生成: [非マッチ, マッチ, 非マッチ, ...] の順に
        segments = []
        last_end = 0
        for ls, le in intersects:
            before = run_text[last_end:ls]
            if before:
                segments.append((before, cur_highlight))
            segments.append((run_text[ls:le], WD_COLOR_INDEX.RED))
            last_end = le
        remaining = run_text[last_end:]
        if remaining:
            segments.append((remaining, cur_highlight))
        if not segments:
            continue

        # 先頭セグメントは既存runを書き換え（テキストとハイライトのみ）
        text0, highlight0 = segments[0]
        set_run_text_preserve_images(run, text0)
        if highlight0 != cur_highlight:
            run.font.highlight_color = highlight0
        # 残りは元runの直後にXMLレベルで挿入し、書式は rpr_template から継承
        prev_r = run._r
        for seg_text, seg_highlight in segments[1:]:
            new_r = OxmlElement('w:r')
            if rpr_template is not None:
                new_r.append(deepcopy(rpr_template))
            t = OxmlElement('w:t')
            t.text = seg_text
            t.set(qn('xml:space'), 'preserve')
            new_r.append(t)
            prev_r.addnext(new_r)
            wrapped = Run(new_r, run._parent)
            wrapped.bold = is_bold
            wrapped.italic = is_italic
            # 元runのハイライトと異なる場合のみ明示設定（同じなら rpr_template の値を尊重）
            if seg_highlight != cur_highlight:
                wrapped.font.highlight_color = seg_highlight
            prev_r = new_r

output_word = "./output/output.docx"
os.makedirs(os.path.dirname(output_word), exist_ok=True)

replacement = {
    '  ':'',       # NBSP + 半角空白 を除去（行冒頭等への混入対策）
    ' ':'',              # 単独NBSPも除去
    '\u0021':'\uFF01',      # !　(全角)
    '\u0022':'\uFF02',      # "　(全角)
    '\u0023':'\uFF03',      # #　(全角)
    '\u0024':'\uFF04',      # $　(全角)
    '\uFF05':'\u0025',      # %　(全角→半角)
    '\uFF06':'\u0026',      # &　(半角) 
    '\u0027':'\uFF07',      # '　(全角)
    '\u0028':'\uFF08',      # (　(全角)
    '\u0029':'\uFF09',      # )　(全角)
    '\uFF0A':'\u002A',      # *　(半角)
    '\u002B':'\uFF0B',      # +　(全角)
    '\uFF0C':'\u002C',      # ,　(半角)　位取りのコンマ
    #'\u002D':'\uFF0D',      # -　(全角)数式記号全角統一だが、ハイフンマイナスだけハイフンとしても使われているため廃止
    '\uFF0E':'\u002E',      # .　(半角)　小数点
    '\uFF10':'\u0030',      # 0　(半角)
    '\uFF11':'\u0031',      # 1　(半角)
    '\uFF12':'\u0032',      # 2　(半角)
    '\uFF13':'\u0033',      # 3　(半角)
    '\uFF14':'\u0034',      # 4　(半角)
    '\uFF15':'\u0035',      # 5　(半角)
    '\uFF16':'\u0036',      # 6　(半角)
    '\uFF17':'\u0037',      # 7　(半角)
    '\uFF18':'\u0038',      # 8　(半角)
    '\uFF19':'\u0039',      # 9　(半角)
    '\u003B':'\uFF1B',      # ;　(全角)
    '\u003C':'\uFF1C',      # <　(全角)
    '\u003D':'\uFF1D',      # =　(全角)
    '\u003E':'\uFF1E',      # >　(全角)
    '\u003F':'\uFF1F',      # ?　(全角)
    '\u0040':'\uFF20',      # @　(全角)
    '\uFF21':'\u0041',      # A　(半角)
    '\uFF22':'\u0042',      # B　(半角)
    '\uFF23':'\u0043',      # C　(半角)
    '\uFF24':'\u0044',      # D　(半角)
    '\uFF25':'\u0045',      # E　(半角)
    '\uFF26':'\u0046',      # F　(半角)
    '\uFF27':'\u0047',      # G　(半角)
    '\uFF28':'\u0048',      # H　(半角)
    '\uFF29':'\u0049',      # I　(半角)
    '\uFF2A':'\u004A',      # J　(半角)
    '\uFF2B':'\u004B',      # K　(半角)
    '\uFF2C':'\u004C',      # L　(半角)
    '\uFF2D':'\u004D',      # M　(半角)
    '\uFF2E':'\u004E',      # N　(半角)
    '\uFF2F':'\u004F',      # O　(半角)
    '\uFF30':'\u0050',      # P　(半角)
    '\uFF31':'\u0051',      # Q　(半角)
    '\uFF32':'\u0052',      # R　(半角)
    '\uFF33':'\u0053',      # S　(半角)
    '\uFF34':'\u0054',      # T　(半角)
    '\uFF35':'\u0055',      # U　(半角)
    '\uFF36':'\u0056',      # V　(半角)
    '\uFF37':'\u0057',      # W　(半角)
    '\uFF38':'\u0058',      # X　(半角)
    '\uFF39':'\u0059',      # Y　(半角)
    '\uFF3A':'\u005A',      # Z　(半角)
    '\u005B':'\uFF3B',      # [　(全角)
    '\u005C':'\uFF3C',      # \　(全角)
    '\u005D':'\uFF3D',      # ]　(全角)
    '\u005E':'\uFF3E',      # ^　(全角)
    '\u005F':'\uFF3F',      # _　(全角)
    '\u0060':'\uFF40',      # `　(全角)
    '\uFF41':'\u0061',      # a　(半角)
    '\uFF42':'\u0062',      # b　(半角)
    '\uFF43':'\u0063',      # c　(半角)
    '\uFF44':'\u0064',      # d　(半角)
    '\uFF45':'\u0065',      # e　(半角)
    '\uFF46':'\u0066',      # f　(半角)
    '\uFF47':'\u0067',      # g　(半角)
    '\uFF48':'\u0068',      # h　(半角)
    '\uFF49':'\u0069',      # i　(半角)
    '\uFF4A':'\u006A',      # j　(半角)
    '\uFF4B':'\u006B',      # k　(半角)
    '\uFF4C':'\u006C',      # l　(半角)
    '\uFF4D':'\u006D',      # m　(半角)
    '\uFF4E':'\u006E',      # n　(半角)
    '\uFF4F':'\u006F',      # o　(半角)
    '\uFF50':'\u0070',      # p　(半角)
    '\uFF51':'\u0071',      # q　(半角)
    '\uFF52':'\u0072',      # r　(半角)
    '\uFF53':'\u0073',      # s　(半角)
    '\uFF54':'\u0074',      # t　(半角)
    '\uFF55':'\u0075',      # u　(半角)
    '\uFF56':'\u0076',      # v　(半角)
    '\uFF57':'\u0077',      # w　(半角)
    '\uFF58':'\u0078',      # x　(半角)
    '\uFF59':'\u0079',      # y　(半角)
    '\uFF5A':'\u007A',      # z　(半角)
    '\u007B':'\uFF5B',      # {　(全角)
    '\u007C':'\uFF5C',      # |　(全角)
    '\u007D':'\uFF5D',      # }　(全角)
    '\u007E':'\uFF5E',      # ~　(全角)
    '\u33A1':'m2',          # 機種依存文字（平方メートル）
    '\u33A5':'m3',          # 機種依存文字（立法メートル）
    '\\ ':'¥',
    '("':'(“',
    '\\':'¥',
    "'":"’",
    '" ':'”\\s',
    ' "':'\\s“',
    '")':'”)',
    '".':'”.',
    '."':'.”',
    '・\t':'• ',
    ' )':')',
    ' . ':'. ',
    ' , ':', ',
    '" ':'” ',
    '( ':'(',
    "' ":"’ ",
    "'":"’",
    "• ":"•   ",
    '） ':'）',
    ' （':'（',
    '."':'.”',
    "' ":"’ ",
    ' 、':'、',
    '。 ':'。',
    ' 年':'年',
    ' 日':'日',
    ' 月':'月',
    '、 ':'、',
    '： ':'：',
    ' ：':'：',
    ' 円':'円',
    ' 時':'時',
    ' 分':'分',
    ' 百万円':'百万円',
    ' 億円':'億円',
    ' ％':'％',
    ' %':'%',
    ',"':',”',
    '  ':' ',
    '",':'”,',
    ' 」':'」',
    #'－':'–',ハイフンマイナス全角のenダッシュ強制変換を廃止
    '（ ':'（',
    ' （':'（',
    '） ':'）',
    '( ':'(',
    ' "':' “',
    '」 ':'」',
    ' 「':'「',
    '］ ':'］',
    'Source：':'Source: ',
    '（Billions）':' (Billions) ',
    'Note：':'Note: ',
    '（Thousands）':'(Thousands) ',
    '（Millions）':' (Millions) ',
    '〜 ':'〜',
    '\t -':'\t\u2013',
    '¥-':'¥–',
    '$ ':'$',
    '➢':'>',
    '\u3000\t':'\t',
    '\\s\\':'\\s¥',
    '\\s\\s\\s':'\\s',
    '\t\\':'\t¥',
}

st.title("テキスト整形（和文）")
st.write("和文フォントの全角・半角への修正や約物の自動変換をします")

option = st.radio("Word文書かテキストか整形対象を選択してください", ("Word文書", "テキスト文書"))

if option == "テキスト文書":
    if "text" not in st.session_state:
        st.session_state.text = ""
    st.session_state.text = st.text_area("テキストを貼りつけてください", height=300, placeholder="ここに貼りつけてください")
    
    if st.button("実行する"):    
        try:
            text = st.session_state.text
            text = normalize_url_email(text)
            for old,new in replacement.items():
                text = text.replace(old, new)
            text = protect_urls(text, apply_slash_colon)
            text = re.sub(r'\s*(\d+)\s*', r'\1', text)
            text = strip_line_head_spaces(text)
        
            st.success("処理が完了しました。下記テキストをコピペしてください")
            with st.container(border=True):
                st.text(text)
            
        except Exception as e:
            st.error(f"ファイルの読み込み中にエラーが発生しました: {e}")
            
else:
    # st.file_uploaderウィジェットの作成
    # type引数で受け付けるファイルの形式を'.docx'に限定
    uploaded_file = st.file_uploader("Wordファイル（.docx）をアップロード", type=['docx'])

    # ファイルがアップロードされた場合の処理
    if uploaded_file is not None:
        file_name = os.path.splitext(uploaded_file.name)[0]  # アップロードされたファイル名を取得
        st.success("ファイルが正常にアップロードされました。")

        # BytesIOでアップロードされたファイルを扱う
        doc_file = BytesIO(uploaded_file.getvalue())

        try:
            # python-docxでWordドキュメントを開く
            doc = docx.Document(doc_file)
            for para in doc.paragraphs:
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)

                # 段落先頭の NBSP(U+00A0) と続く空白を除去（最初のテキストrunに対して）
                for r in para.runs:
                    if r.text:
                        stripped = re.sub('^ [  ]*', '', r.text)
                        if stripped != r.text:
                            set_run_text_preserve_images(r, stripped)
                        break

                # Run単位の処理：テキスト置換 + 数字前後空白除去（画像は保持）
                for run in para.runs:
                    t = run.text
                    for old, new in replacement.items():
                        t = t.replace(old, new)
                    t = strip_spaces_around_numbers(t)
                    set_run_text_preserve_images(run, t)

                # 段落全体でURL正規化 + スラッシュ・コロン変換（URL跨ぎrun対応）
                if para.runs:
                    full_text = ''.join(run.text for run in para.runs)
                    full_text = normalize_url_email(full_text)
                    full_text = protect_urls(full_text, apply_slash_colon)
                    full_text = strip_line_head_spaces(full_text)
                    idx = 0
                    for run in para.runs:
                        run_len = len(run.text)
                        set_run_text_preserve_images(run, full_text[idx:idx + run_len])
                        idx += run_len

                # 行頭（段落先頭・改行後）の先頭空白（NBSP・半角スペース）を除去
                _at_line_start = True
                for _elem in para._element.iter():
                    if _elem.tag == qn('w:br'):
                        _at_line_start = True
                    elif _elem.tag == qn('w:t') and _elem.text:
                        if _at_line_start:
                            _elem.text = _elem.text.lstrip('  	\u3000')
                        if _elem.text:
                            _at_line_start = False

                # Run単位のボールド・斜体・ダッシュハイライト
                apply_run_highlights(para)

                # ※/*+数字パターンのハイライト（run単位で処理）
                apply_note_highlight_per_run(para)

                # 箇条書きスタイルの解除とビュレット挿入
                is_list = para.style.name.startswith('List') or para._element.pPr is not None and para._element.pPr.find(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr') is not None
                if is_list:
                    if para.runs:
                        set_run_text_preserve_images(para.runs[0], '• ' + para.runs[0].text)
                    # numPr（ナンバリング属性）を削除して箇条書き設定を解除
                    if para._element.pPr is not None:
                        numPr = para._element.pPr.find(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                        if numPr is not None:
                            para._element.pPr.remove(numPr)
                    para.style = doc.styles['Normal']

            # 変更を新しいPDFファイルに保存
            # output_word = "./output/output.docx"
            doc.save(output_word)
            st.success("処理が完了しました")

            # ドキュメントの内容を表示
            #st.subheader("処理されたWordファイルの内容")
        
            # ドキュメント内の各段落を読み込んで表示
            #for para in doc.paragraphs:
            #    st.write(para.text)
            st.success("ダウンロードボタンを押してください")
            with open(output_word, "rb") as file:
                word_data = file.read()
                # ダウンロードボタンを作成
                st.download_button(
                    label="Word文書をダウンロード",
                    data=word_data,
                    file_name=file_name+"_chk.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", # WordファイルのMIMEタイプ
                    on_click="ignore" # 再実行を無視する設定
                )
            print(f"ダウンロードしました。")
        
        except Exception as e:
            st.error(f"ファイルの読み込み中にエラーが発生しました: {e}")
    else:
        st.info("ファイルをアップロードしてください。")
            
